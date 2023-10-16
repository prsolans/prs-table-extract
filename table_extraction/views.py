""" views.py """
import json
import uuid
from django.shortcuts import render
from docx import Document as DocxDocument
from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import generics, status
from table_extraction.forms import DocumentForm
from .models import Document
from .serializers import TableDataSerializer, DocumentSerializer

def index(request):
    context = {
        "title": "Django example",
    }
    return render(request, "index.html", context)

def parse_word_document(file_path):
    doc = DocxDocument(file_path)
    return doc

def extract_table_data(doc):
    tables_data = []

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text)
            table_data.append(row_data)
        tables_data.append(table_data)

    return tables_data

def upload_doc(request):
    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)

        if form.is_valid():
            instance = form.save()  # Save the uploaded file instance
            doc = DocxDocument(instance.upload.path)

            tables_data = extract_table_data(doc)

            json_string_1 = json.dumps(tables_data[0], indent=4)
            json_string_2 = json.dumps(tables_data[1], indent=4)

            # Do something with tables_data, e.g., send it to a template, save it to DB, etc.
            return render(request, 'results.html', {'data_table': json_string_1, 'data_table2' : json_string_2})
    else:
        form = DocumentForm()

    return render(request, 'upload.html', {'form': form})

def results(request):
    context = {
        "title": "Data Results",
    }
    return render(request, "results.html", context)

@api_view(['POST'])
def table_extraction_api(request):
    if request.method == 'POST':
        word_file = request.FILES.get('document')
        table_data = upload_doc(word_file)  # your logic to extract tables
        serializer = TableDataSerializer(data={'data': table_data})

        if serializer.is_valid():
            return Response(serializer.data)
        else:
            return Response(serializer.errors, status=400)

class DocumentAPI(generics.CreateAPIView):
    queryset = Document.objects.all()
    serializer_class = DocumentSerializer

    def create(self, request, *args, **kwargs):
        # This uses the default creation logic provided by DRF
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        # Once validated, save the object and parse the document
        instance = serializer.save()
        doc = parse_word_document(instance.upload.path)
        tables_data = extract_table_data(doc)
        # Return the parsed tables_data as a response
        return Response(tables_data, status=status.HTTP_201_CREATED)

class GuidAPI(generics.GenericAPIView):
    def post(self, request, *args, **kwargs):
        received_guid = request.data.get('guid', None)
        if not received_guid:
            return Response({'detail': 'GUID is required.'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            valid_uuid = uuid.UUID(received_guid, version=4)
        except ValueError:
            # If it's a value error, then the string is not a valid hex code for a UUID
            return Response({'detail': 'Invalid GUID provided.'}, status=status.HTTP_400_BAD_REQUEST)

        # Now you have a valid UUID in the variable `valid_uuid` and can process it as needed
        # For now, let's just return it in the response.
        return Response({'received_guid': str(valid_uuid)}, status=status.HTTP_200_OK)